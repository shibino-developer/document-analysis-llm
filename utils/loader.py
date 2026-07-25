from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader


class DocumentLoader:
    """
    Loads PDF, DOCX and TXT files and returns
    LangChain Document objects.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def load(self, uploaded_file) -> List[Document]:
        """
        Detect file type and load the document.

        Parameters
        ----------
        uploaded_file : UploadedFile
            Streamlit UploadedFile

        Returns
        -------
        List[Document]
        """

        if uploaded_file is None:
            raise ValueError("No file uploaded.")

        extension = Path(uploaded_file.name).suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}"
            )

        if extension == ".pdf":
            return self._load_pdf(uploaded_file)

        if extension == ".docx":
            return self._load_docx(uploaded_file)

        if extension == ".txt":
            return self._load_txt(uploaded_file)

        raise ValueError("Unsupported document.")

    # ---------------------------------------------------------
    # PDF
    # ---------------------------------------------------------

    def _load_pdf(self, uploaded_file) -> List[Document]:

        reader = PdfReader(uploaded_file)

        documents = []

        total_pages = len(reader.pages)

        for page_number, page in enumerate(reader.pages, start=1):

            text = page.extract_text()

            if not text:
                continue

            documents.append(

                Document(

                    page_content=text,

                    metadata={

                        "source": uploaded_file.name,

                        "file_type": "pdf",

                        "page": page_number,

                        "total_pages": total_pages

                    }

                )

            )

        return documents

    # ---------------------------------------------------------
    # DOCX
    # ---------------------------------------------------------

    def _load_docx(self, uploaded_file) -> List[Document]:

        doc = DocxDocument(uploaded_file)

        paragraphs = []

        for para in doc.paragraphs:

            if para.text.strip():

                paragraphs.append(para.text)

        text = "\n".join(paragraphs)

        return [

            Document(

                page_content=text,

                metadata={

                    "source": uploaded_file.name,

                    "file_type": "docx",

                    "paragraphs": len(paragraphs)

                }

            )

        ]

    # ---------------------------------------------------------
    # TXT
    # ---------------------------------------------------------

    def _load_txt(self, uploaded_file) -> List[Document]:

        text = uploaded_file.read().decode("utf-8")

        return [

            Document(

                page_content=text,

                metadata={

                    "source": uploaded_file.name,

                    "file_type": "txt",

                    "lines": len(text.splitlines())

                }

            )

        ]